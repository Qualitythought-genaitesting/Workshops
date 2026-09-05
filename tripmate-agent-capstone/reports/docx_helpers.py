"""Small python-docx helpers for consistently branded Quality Thought documents."""
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

NAVY = RGBColor(0x1F, 0x38, 0x64)
ORANGE = RGBColor(0xEE, 0x4C, 0x12)
MUTED = RGBColor(0x64, 0x74, 0x8B)
FONT = "Calibri"


def new_document(title: str, subtitle: str, meta: dict) -> Document:
    doc = Document()
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(0.9)
        s.top_margin = s.bottom_margin = Inches(0.8)
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(10.5)
    for lvl, size in ((1, 18), (2, 14), (3, 12)):
        h = doc.styles[f"Heading {lvl}"]
        h.font.name = FONT; h.font.size = Pt(size); h.font.bold = True; h.font.color.rgb = NAVY
    # cover block
    p = doc.add_paragraph(); r = p.add_run("QUALITY THOUGHT  ·  AI AGENT TESTING CAPSTONE"); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = ORANGE
    p = doc.add_paragraph(); r = p.add_run(title); r.bold = True; r.font.size = Pt(26); r.font.color.rgb = NAVY
    p = doc.add_paragraph(); r = p.add_run(subtitle); r.font.size = Pt(12); r.font.color.rgb = MUTED
    t = doc.add_table(rows=0, cols=2); t.style = "Table Grid"
    for k, v in meta.items():
        row = t.add_row().cells
        row[0].text, row[1].text = k, str(v)
        row[0].paragraphs[0].runs[0].bold = True
        shade(row[0], "F3F5F9")
    set_col_widths(t, [1.8, 4.9])
    doc.add_paragraph()
    return doc


def shade(cell, hex_fill: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def set_col_widths(table, widths_in):
    table.autofit = False
    for row in table.rows:
        for i, w in enumerate(widths_in):
            if i < len(row.cells):
                row.cells[i].width = Inches(w)


def h(doc, text, level=1):
    return doc.add_heading(text, level=level)


def para(doc, text, bold=False, italic=False, color=None, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    if size: r.font.size = Pt(size)
    return p


def bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        if isinstance(it, tuple):
            r = p.add_run(it[0] + " "); r.bold = True
            p.add_run(it[1])
        else:
            p.add_run(it)


def numbered(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Number")


def table(doc, headers, rows, widths=None, font_size=9, header_fill="1F3864"):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hd in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        r = c.paragraphs[0].add_run(hd); r.bold = True; r.font.size = Pt(font_size); r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade(c, header_fill)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            txt = str(v)
            r = cells[i].paragraphs[0].add_run(txt); r.font.size = Pt(font_size)
            if txt in ("PASS", "Pass"):
                r.bold = True; r.font.color.rgb = RGBColor(0x0B, 0x6B, 0x3A)
            elif txt in ("FAIL", "Fail"):
                r.bold = True; r.font.color.rgb = RGBColor(0x92, 0x2B, 0x21)
            elif txt == "Flaky":
                r.bold = True; r.font.color.rgb = RGBColor(0x9C, 0x4A, 0x00)
            if ri % 2:
                shade(cells[i], "F8FAFC")
    if widths:
        set_col_widths(t, widths)
    doc.add_paragraph()
    return t


def callout(doc, text, fill="FFF4EC"):
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"
    c = t.rows[0].cells[0]; c.text = ""
    r = c.paragraphs[0].add_run(text); r.font.size = Pt(10); r.italic = True
    shade(c, fill)
    doc.add_paragraph()
